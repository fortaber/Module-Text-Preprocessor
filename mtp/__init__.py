# import spacy
# from spacy.matcher import Matcher
import docx
import pathlib
# import mpire
import os
# from mtp.config.students_keyWords import STUDENTS_IDENTIFICATIONS_WORDS
# from mtp.config.teachers_keyWords import TEACHERS_IDENTIFICATIONS_WORDS
# import typing


# def extract_names(text):
#     doc = nlp(text)
#     return [ent.text for ent in doc.ents if ent.label_ == "PER"]

# def lookingup(parallel: bool, array: list):
#     """
#     поиск статуса ФИО\n
#     parallel[bool] - распараллеливание поиска\n
#     array[list] - лист с нормализованным текстом\n
#     id_fio[int] - id расположения фио 
#     """
#     if not parallel:


# !!! не работает !!!
# def parallel_search_fio(p_text: str):

#     """
#     параллельная реализация функции search_fio для поиска имён в тексте
#     """

#     result = {
#         "st": [],
#         "tc": []
#     }

#     print(p_text)
#     full_name_st = mdr.search_fio(p_text, True)
#     if full_name_st is not None:
#         result['st'].append(full_name_st)

#     full_name_tc = mdr.search_fio(p_text, False)
#     if full_name_tc is not None:
#         result["tc"].append(full_name_tc)

#     return result

class mtp:
    def __init__(self):
        
        """
        Создание объекта Module Text Preprocessor
        """
        self.path = "" # temp
        
        self.raw_text = ""
        self.clean_text = ""

        try:
            # self.nlp = spacy.load("ru_core_news_lg")
            print("Успешная инициализация НЛП модели")

        except OSError:
            print("Ошибка: spaCy модель 'ru_core_news_lg' не была найдена.")
            print("Возможное решение: python -m spacy download ru_core_news_lg")

    def __call__(self):
        return self

    def read_document(self, path_: str):
        """
        добавление и считывания файла .docx формата\n
        **path_** - путь к файлу .docx
        """
        if '/' not in path_:
            self.path = f"{pathlib.Path(__file__).parent.resolve()}{self.path}"
        self.path = path_

        try:
            doc = docx.Document(self.path)

            fullText = []

            # чтение текста
            for p in doc.paragraphs:
                fullText.append(p.text)

            self.raw_text = '\n'.join(fullText)
            
            # чтение таблиц
            for table in doc.tables:
                table_ = []
                for row in table.rows:
                    row_ = []
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            row_.append(p.text)
                    
                    row_ = ' '.join(row_)
                    table_.append(row_)
            # print(self.tables_)
                      
        except:
            print("ERROR: Не удалось найти файл!")



    # def set_cpus(self, num: int):
    #     """
    #     установление количество процессоров, на которое будет распараллелен поиск ФИО студентов и преподавателей
    #     **num** - количество процессоров для распараллеливания через mpire
    #     """
    #     if num <= 0 or num > os.cpu_count():
    #         print("Некорректное количество процессоров")
    #         return
    #     self.cpus = num


    
    # def set_n(self, n: int):
    #     """
    #     Установление количества ФИО спанов, которые будут добавляться с конца и начала документа .docx\n
    #     **n** - количество найденных spaCy НЛП ФИО с начала документа и конца, которые берутся на дальнейшую обработку
    #     """

    #     self.n = n

    # def get_names(self) -> list:
    #     """
    #     Получение всех имён с тегами PER, которые удалось найти NLP модели
    #     """
    #     return self.names_


# +++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++ Moe
    # def search_fio(self, text_: str, st: bool) -> list:
    #     """
    #     поиск ФИО студентов/преподавателей\n
    #     **text_** - строка для поиска\n
    #     **st** (True) - поиск студетов
    #     """

    #     pattern = []

    #     if st:
    #         pattern = [
    #         {"LEMMA": { "IN": STUDENTS_IDENTIFICATIONS_WORDS}},
    #         {"IS_PUNCT":True, "OP": "*"},
    #         {"OP": "*"},
    #         {"TEXT": {"REGEX": "^[А-ЯЁа-яё][а-яё-]+$"}}
    #     ]
    #     else:
    #         pattern = [
    #         {"LEMMA": { "IN": TEACHERS_IDENTIFICATIONS_WORDS}},
    #         {"IS_PUNCT":True, "OP": "*"},
    #         {"OP": "*"},
    #         {"TEXT": {"REGEX": "^[А-ЯЁа-яё][а-яё-]+$"}}
    #     ]

    #     doc = self.nlp(text_)
    #     matcher = Matcher(self.nlp.vocab)
    #     matcher.add("full_names_st", [pattern])
    #     matches = matcher(doc)

    #     res_ = []
        
    #     for match_id, start, end in matches:
            
    #         span = doc[start:end]

    #         res_.append(span)
        
    #     if len(res_) == 0:
    #         return

    #     for i in range(len(res_) - 2, -1, -1):
    #         if res_[i].text in res_[i + 1].text:
    #             del res_[i]
        
    #     # print(f"res2: {res_}")

    #     tokens = []

    #     for span_ in res_:
    #         for token in span_:
    #             if token.ent_type_ == "PER":
    #                 tokens.append(token)

    #     if len(tokens) < 1:
    #         return
        
    #     # print(f"tokens: {tokens}\n")
    #     return tokens         

    def extract_names(self):
        
        """
        Нахождение имен в при помощи NLP библиотеки spaCy
        """

        if not self.nlp or not self.text_:
            print("Не удалось найти имена: модель не была загружена или документ пустой")
            return []

        doc = self.nlp(self.text_)

        # проверка на то, чтобы n всегда было в range для doc.ents.count 
        if self.n > doc.ents.__len__():
            self.n = doc.ents.__len__()
            print(f"doc.ents.__len__() = {doc.ents.__len__()}")

        # !!! старый метод поиска имён в тексте - плохо находит и распазнаёт "Рис", как ent с label = "PERSON"
        # поиск ФИО в тексте (глобально)
        # name = []
        # cnt = 0
        # while cnt < self.n:
        #     for i in range(len(doc.ents)):
        #         if doc.ents[i].label_ == "PER":
        #             # сохранение строки и её расположения
        #             # self.names_.append([doc.ents[i].text, [doc.ents[i].start_char, doc.ents[i].end_char]])
        #             name.append(doc.ents[i].text)
        #             print(doc.ents[i].text)
        #         else:
        #             if len(name) > 0:
        #                 self.names_.append(name)
        #                 print(name)
        #                 cnt += 1
        #                 name = []
        #         # self.names_.append([doc.ents[i].text])
        #     # if doc.ents[-i - 1].label == "PER":
        #         # self.names_.append([doc.ents[-i - 1].text])

        # поиск имён в таблице
        # !!! изначально предполагается, что ФИО авторов будут находится в таблице !!!
        for i in range(len(self.tables_)):

            text_ = self.tables_[i]

            if isinstance(self.tables_[i], list):
                for j in range(len(text_)):
                    for st in [True, False]:
                        full_name = self.search_fio(text_[j], st)

                        if full_name is not None:
                            if st:
                                self.students.append(full_name)
                            else:
                                self.teachers.append(full_name)

            else:
                for st in [True, False]:
                    full_name = self.search_fio(text_, st)
                    
                    if full_name is not None:
                        if st:
                            self.students.append(full_name)
                        else:
                            self.teachers.append(full_name)

        # поиск во всём тексте
        try:
            doc = docx.Document(self.path)

            for p_text in doc.paragraphs:
                # print(p_text.text)
                for st in [True, False]:
                        full_name = self.search_fio(p_text.text, st)

                        if full_name is not None:
                            if st:
                                self.students.append(full_name)
                            else:
                                self.teachers.append(full_name)

        #     # !!! не работает !!!
        #     # paragraphes = [p.text for p in doc.paragraphs]
        #     # if __name__ == "__main__":
        #     #     with mpire.WorkerPool(n_jobs=self.cpus) as pool:
        #     #         for res in pool.map(
        #     #             parallel_search_fio,
        #     #             paragraphes
        #     #     ):
        #     #             self.students.extend(res["students"])
        #     #             self.teachers.extend(res["teachers"])
        
        #     # for st in [True, False]:
        #     #     full_name = self.search_fio(self.text_, st)

        #     #     if full_name is not None:
        #     #         if st:
        #     #             self.students.append(full_name)
        #     #         else:
        #     #             self.teachers.append(full_name)
        except:
            print("Ошибка чтения документа")
            return

        print("Студенты")
        print(self.students)
        print("\nПреподаватели")
        print(self.teachers)
        # print(self.names_)

            # doc = self.nlp(self.tables_[i][0])
            # matcher = Matcher(self.nlp.vocab)
            # matcher.add("full_names_st", [patterns_s])
            # matches = matcher(doc)

            # print(f"len(matches) = {len(matches)}")

            # res_ = []
            # print(doc.ents.__len__())
            
            # for match_id, start, end in matches:
                
            #     span = doc[start:end]
            #     print(f"span: {span}")

            #     res_.append(span.text)
            
            # for i in range(len(res_) - 2, -1, -1):
            #     if res_[i] in res_[i + 1]:
            #         del res_[i]

            # print(res_)

        # print(f"Удалось обнаружить {self.names_.__len__()} предположительных имён в тексте и {self.names_tables_.__len__()} в таблицах")
        # print(self.names_)